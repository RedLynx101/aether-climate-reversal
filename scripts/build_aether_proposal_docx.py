"""Build the AETHER proposal DOCX and PDF from manuscript/proposal/aether_proposal_source.md.

Run with: uv run --with python-docx,docx2pdf,pymupdf python scripts/build_aether_proposal_docx.py
"""

from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from publication_safety import resolve_within, trusted_windows_powershell

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "proposal" / "aether_proposal_source.md"
FIG_DIR = ROOT / "analysis" / "figures"
FIG_CACHE = ROOT / "manuscript" / "proposal" / "figcache"
OUT_DIR = ROOT / "manuscript" / "proposal"
DOCX_OUT = OUT_DIR / "AETHER_Conditional_Feasibility_Proposal.docx"
PDF_OUT = OUT_DIR / "AETHER_Conditional_Feasibility_Proposal.pdf"
WORD_PDF_EXPORTER = ROOT / "scripts" / "export_aether_proposal_pdf.ps1"

BODY_FONT = "Cambria"
HEAD_FONT = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)
HEAD_INK = RGBColor(0x10, 0x2A, 0x43)
MUTED = RGBColor(0x5B, 0x67, 0x7A)
HEADER_FILL = "EAF0F6"
ALT_FILL = "F5F8FB"
BORDER = "B9C4D0"


def set_font(run, name=BODY_FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(par, text, size, color=INK, base_bold=False, font=BODY_FONT):
    """Add text with **bold** inline markup."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not chunk:
            continue
        bold = base_bold or (i % 2 == 1)
        run = par.add_run(chunk)
        set_font(run, name=font, size=size, color=color, bold=bold)


def keep_with_next(par):
    ppr = par._p.get_or_add_pPr()
    node = OxmlElement("w:keepNext")
    ppr.append(node)


def keep_together(par):
    """Prevent a paragraph from splitting across pages when Word can fit it."""
    ppr = par._p.get_or_add_pPr()
    node = OxmlElement("w:keepLines")
    ppr.append(node)

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), BORDER)
        borders.append(el)
    tbl_pr.append(borders)


def strip_source_footer(path: Path) -> Path:
    """Remove the bottom-most isolated text band (the 'Source: ...' footer)
    from a figure PNG when it sits in the bottom 14% of the image."""
    from PIL import Image

    FIG_CACHE.mkdir(exist_ok=True)
    out = FIG_CACHE / path.name
    img = Image.open(path)
    gray = img.convert("L")
    w, h = gray.size
    px = gray.load()
    dark_rows = []
    for y in range(h):
        dark = 0
        for x in range(0, w, 3):
            if px[x, y] < 185:
                dark += 1
                if dark >= 3:
                    dark_rows.append(y)
                    break
    # group contiguous dark rows (gap tolerance 6px to absorb antialiasing)
    bands = []
    for y in dark_rows:
        if bands and y - bands[-1][1] <= 6:
            bands[-1][1] = y
        else:
            bands.append([y, y])
    if len(bands) >= 2:
        last = bands[-1]
        prev = bands[-2]
        band_h = last[1] - last[0] + 1
        gap = last[0] - prev[1]
        if last[0] > 0.86 * h and band_h < 40 and gap >= 8:
            crop_y = max(prev[1] + 2, last[0] - 4)
            img = img.crop((0, 0, w, crop_y))
    img.save(out)
    return out


def add_page_number_footer(section):
    footer = section.footer
    par = footer.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    set_font(run, name=HEAD_FONT, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)


class Builder:
    def __init__(self):
        self.doc = Document()
        self.fig_n = 0
        self.tbl_n = 0
        sec = self.doc.sections[0]
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11)
        for attr in ("left_margin", "right_margin"):
            setattr(sec, attr, Inches(1.0))
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        add_page_number_footer(sec)
        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal.font.size = Pt(10.5)

    def para(self, text, size=10.5, color=INK, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=8, bold=False, italic=False, font=BODY_FONT,
             indent=None, line=1.12):
        par = self.doc.add_paragraph()
        par.alignment = align
        pf = par.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = line
        if indent is not None:
            pf.left_indent = Inches(indent)
        add_inline(par, text, size=size, color=color, base_bold=bold, font=font)
        if italic:
            for run in par.runs:
                run.italic = True
        keep_together(par)
        return par

    def heading(self, text, level):
        sizes = {1: 14.5, 2: 12, 3: 11}
        before = {1: 18, 2: 14, 3: 10}
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = par.paragraph_format
        pf.space_before = Pt(before[level])
        pf.space_after = Pt(6)
        pf.line_spacing = 1.05
        add_inline(par, text, size=sizes[level], color=HEAD_INK, base_bold=True, font=HEAD_FONT)
        keep_with_next(par)
        return par

    def quote(self, text):
        par = self.para(text, size=10.5, color=HEAD_INK, align=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=4, space_after=10, indent=0.4)
        for run in par.runs:
            run.italic = True
        return par

    def bullet(self, text):
        par = self.doc.add_paragraph(style="List Bullet")
        pf = par.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(5)
        pf.line_spacing = 1.12
        pf.left_indent = Inches(0.3)
        add_inline(par, text, size=10.5, color=INK)
        keep_together(par)
        return par

    def figure(self, filename, caption):
        path = resolve_within(FIG_DIR, filename, FIG_DIR)
        if not path.exists():
            raise FileNotFoundError(path)
        path = strip_source_footer(path)
        self.fig_n += 1
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_before = Pt(10)
        par.paragraph_format.space_after = Pt(2)
        run = par.add_run()
        run.add_picture(str(path), width=Inches(5.9))
        keep_with_next(par)
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        cap.paragraph_format.line_spacing = 1.05
        add_inline(cap, f"**Figure {self.fig_n}.** {caption}", size=9, color=MUTED, font=HEAD_FONT)

    def table(self, rows, caption=None):
        self.tbl_n += 1
        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cap.paragraph_format.space_before = Pt(10)
            cap.paragraph_format.space_after = Pt(4)
            add_inline(cap, f"**Table {self.tbl_n}.** {caption}", size=9, color=MUTED, font=HEAD_FONT)
            keep_with_next(cap)
        n_cols = len(rows[0])
        font_size = 9 if n_cols <= 6 else 7.5
        table = self.doc.add_table(rows=len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_borders(table)
        for r, row in enumerate(rows):
            tr_pr = table.rows[r]._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
            if r == 0:
                tbl_header = OxmlElement("w:tblHeader")
                tr_pr.append(tbl_header)
            for c, text in enumerate(row):
                cell = table.cell(r, c)
                cell.paragraphs[0].text = ""
                par = cell.paragraphs[0]
                par.paragraph_format.space_before = Pt(2)
                par.paragraph_format.space_after = Pt(2)
                par.paragraph_format.line_spacing = 1.0
                if r == 0:
                    set_cell_shading(cell, HEADER_FILL)
                    add_inline(par, text, size=font_size, color=HEAD_INK, base_bold=True, font=HEAD_FONT)
                else:
                    if r % 2 == 0:
                        set_cell_shading(cell, ALT_FILL)
                    add_inline(par, text, size=font_size, color=INK, font=HEAD_FONT)
        # Keep compact tables whole when possible. For longer tables, keep the
        # repeated header with at least the first data row.
        keep_row_count = max(0, len(rows) - 1) if len(rows) <= 7 else min(2, len(rows))
        for keep_row in range(keep_row_count):
            for keep_cell in table.rows[keep_row].cells:
                for keep_par in keep_cell.paragraphs:
                    keep_with_next(keep_par)
                    keep_together(keep_par)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        spacer.paragraph_format.line_spacing = 1.0
        return table

    def reference(self, text):
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = par.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = 1.05
        pf.left_indent = Inches(0.35)
        pf.first_line_indent = Inches(-0.35)
        add_inline(par, text, size=9.5, color=INK)
        keep_together(par)
        return par

    def title_page(self, meta):
        for _ in range(6):
            self.doc.add_paragraph()
        t = self.para(meta["TITLE"], size=24, color=HEAD_INK, align=WD_ALIGN_PARAGRAPH.CENTER,
                      bold=True, font=HEAD_FONT, space_after=14, line=1.1)
        self.para(meta["SUBTITLE"], size=13.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
                  font=HEAD_FONT, space_after=40, line=1.25)
        self.para(meta["AUTHOR"], size=13, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER,
                  bold=True, font=HEAD_FONT, space_after=6)
        self.para(meta["DATE"], size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
                  font=HEAD_FONT, space_after=4)
        self.para(meta["STATUS"], size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER,
                  font=HEAD_FONT, space_after=0)
        if meta.get("REPOSITORY"):
            self.para(meta["REPOSITORY"], size=9.5, color=MUTED,
                      align=WD_ALIGN_PARAGRAPH.CENTER, font=HEAD_FONT,
                      space_before=8, space_after=0)
        self.doc.add_page_break()


def parse_and_build():
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    b = Builder()
    meta = {}
    i = 0
    # strip leading html comment
    while i < len(lines) and not lines[i].startswith("TITLE:"):
        i += 1
    while i < len(lines) and ":" in lines[i] and lines[i].split(":")[0] in (
            "TITLE", "SUBTITLE", "AUTHOR", "DATE", "STATUS", "REPOSITORY"):
        key, val = lines[i].split(":", 1)
        meta[key] = val.strip()
        i += 1
    b.title_page(meta)

    pending_caption = None
    in_references = False
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("<!--"):
            i += 1
            continue
        if line.startswith("#### "):
            b.heading(line[5:], 3)
        elif line.startswith("### "):
            b.heading(line[4:], 3)
        elif line.startswith("## "):
            text = line[3:]
            in_references = text.strip().lower() == "references"
            if in_references:
                b.doc.add_page_break()
            b.heading(text, 1)
        elif line.startswith("TABLE-CAPTION:"):
            pending_caption = line.split(":", 1)[1].strip()
        elif line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw = lines[i].strip().strip("|")
                cells = [c.strip() for c in raw.split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            b.table(rows, caption=pending_caption)
            pending_caption = None
            continue
        elif line.startswith("FIG: "):
            fname, caption = [p.strip() for p in line[5:].split("|", 1)]
            b.figure(fname, caption)
        elif line.startswith("> "):
            b.quote(line[2:])
        elif line.startswith("- "):
            if in_references:
                b.reference(line[2:])
            else:
                b.bullet(line[2:])
        else:
            b.para(line)
        i += 1

    b.doc.save(DOCX_OUT)
    print(f"Saved {DOCX_OUT}")
    print(f"Figures: {b.fig_n}, Tables: {b.tbl_n}")


def export_pdf():
    try:
        from docx2pdf import convert
    except ModuleNotFoundError:
        subprocess.run(
            [
                str(trusted_windows_powershell()),
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(WORD_PDF_EXPORTER),
                "-DocxPath",
                str(DOCX_OUT),
                "-PdfPath",
                str(PDF_OUT),
            ],
            check=True,
        )
    else:
        convert(str(DOCX_OUT), str(PDF_OUT))
    print(f"Saved {PDF_OUT}")


def render_preview(pages=None):
    try:
        import fitz
    except ModuleNotFoundError:
        print("PyMuPDF is unavailable; preview rendering skipped.")
        return
    out = ROOT / "manuscript" / "proposal" / "preview"
    out.mkdir(exist_ok=True)
    for old in out.glob("page_*.png"):
        old.unlink()
    pdf = fitz.open(str(PDF_OUT))
    print(f"PDF pages: {pdf.page_count}")
    for idx in range(pdf.page_count):
        if pages and (idx + 1) not in pages:
            continue
        page = pdf.load_page(idx)
        pix = page.get_pixmap(dpi=100)
        pix.save(str(out / f"page_{idx + 1:02d}.png"))
    print(f"Previews in {out}")


if __name__ == "__main__":
    parse_and_build()
    if "--no-pdf" not in sys.argv:
        export_pdf()
        render_preview()

