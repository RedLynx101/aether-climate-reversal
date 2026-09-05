from __future__ import annotations

import re
import sys
from pathlib import Path

# This v0.45 builder writes obsolete submission metadata. Keep it only for
# explicit historical recovery, and fail before optional legacy dependencies
# or any mutation are reached during normal use.
CANONICAL_ROOT = Path(__file__).resolve().parents[1]
if "--legacy-v0-45-rebuild" not in sys.argv[1:]:
    raise SystemExit(
        "Retired v0.45 Word builder. Use `python scripts/build_current_publication.py` "
        "or `--check`; historical recovery requires --legacy-v0-45-rebuild."
    )
try:
    legacy_target = Path(sys.argv[sys.argv.index("--legacy-output-dir") + 1]).resolve()
except (ValueError, IndexError):
    raise SystemExit(
        "Legacy recovery requires --legacy-output-dir PATH naming a separate, "
        "complete isolated AETHER checkout. The active checkout is never a legacy target."
    )
if legacy_target == CANONICAL_ROOT or not (legacy_target / "pyproject.toml").is_file():
    raise SystemExit("--legacy-output-dir must be a complete AETHER checkout distinct from the active checkout.")
LEGACY_ROOT = legacy_target

import zipfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from publication_safety import resolve_within

try:
    from PIL import Image
except Exception:  # pragma: no cover - render-time dependency check
    Image = None


ROOT = LEGACY_ROOT
PAPER = ROOT / "manuscript" / "paper" / "aether_scientific_paper.md"
OUTPUT = ROOT / "manuscript" / "submission" / "AETHER_Atmospheric_Engineering_Through_High_Energy_Removal_v0.45.docx"
MANIFEST = ROOT / "manuscript" / "submission" / "aether_word_export_manifest.md"
NOTES = ROOT / "research" / "parameters" / "word-export-and-docx-readiness-notes.md"
FIGURE_DIR = ROOT / "analysis" / "figures"
VERSION = "v0.45"
DOCX_NAME = OUTPUT.name
REPOSITORY_URL = "https://github.com/RedLynx101/aether-climate-reversal"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = RGBColor(232, 238, 245)
LIGHT_GRAY = RGBColor(242, 244, 247)
MUTED = RGBColor(91, 103, 122)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, name: str = "Calibri", size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D9E1EA", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_row_repeat_as_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def keep_compact_table_together(table, data_row_count: int) -> None:
    if data_row_count > 6:
        return
    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


def paragraph_border_bottom(paragraph, color: str = "B7C4D4", size: str = "8") -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    if "AETHER Caption" not in styles:
        caption = styles.add_style("AETHER Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["AETHER Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)

    if "AETHER Table Text" not in styles:
        table_text = styles.add_style("AETHER Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = styles["AETHER Table Text"]
    table_text.font.name = "Calibri"
    table_text._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    table_text._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    table_text.font.size = Pt(8.7)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.12

    if "AETHER Callout" not in styles:
        callout = styles.add_style("AETHER Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["AETHER Callout"]
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    callout._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = NAVY
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run(f"AETHER | Working paper {VERSION}")
    set_run_font(hr, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(fp.add_run("Page "), size=9, color=MUTED)
    add_field(fp, "PAGE")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("AETHER RESEARCH MANUSCRIPT")
    set_run_font(r, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    r = title.add_run("AETHER")
    set_run_font(r, size=34, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("Atmospheric Engineering Through High-Energy Removal")
    set_run_font(r, size=17, color=BLUE, bold=True)

    deck = doc.add_paragraph()
    deck.paragraph_format.space_after = Pt(16)
    r = deck.add_run("A Conditional Feasibility Analysis for 100 GtCO2/year Removal in an AI- and Robotics-Accelerated Economy")
    set_run_font(r, size=13, color=MUTED, italic=True)

    rule = doc.add_paragraph()
    paragraph_border_bottom(rule)

    rows = [
        ("Prepared by", "Noah Hicks"),
        ("Document type", "Working scientific paper and academic review package"),
        ("Version", VERSION),
        ("Generated", "August 10, 2026"),
        ("Repository", REPOSITORY_URL),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 9360)
    set_table_borders(table, "FFFFFF", "0")
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        for cell, width in zip(cells, [1700, 7660]):
            set_cell_width(cell, width)
            set_cell_margins(cell, top=55, bottom=55, start=0, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles["AETHER Table Text"]
            p.paragraph_format.space_after = Pt(0)
        lr = cells[0].paragraphs[0].add_run(label.upper())
        set_run_font(lr, size=8.5, color=MUTED, bold=True)
        vr = cells[1].paragraphs[0].add_run(value)
        set_run_font(vr, size=10.5, color=BLACK)

    doc.add_paragraph()
    callout_table = doc.add_table(rows=1, cols=1)
    set_table_width(callout_table, 9360)
    set_table_borders(callout_table, "CAD5E2", "6")
    cell = callout_table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.style = doc.styles["AETHER Callout"]
    r = p.add_run("Review posture: ")
    set_run_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run("This is a serious proposal and model package, not a journal-ready climate result. The current failing gates are climate-model publication quality and species-emissions inputs; those remain visible by design.")
    set_run_font(r, size=10.5, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Built from the {VERSION} AETHER repository submission manuscript. Figures, citations, evidence classes, and unresolved gates are preserved for academic review.")
    set_run_font(r, size=10, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_runs_from_markdown(paragraph, text: str, base_size: float | None = None) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=base_size or 9.5, color=NAVY)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=base_size, bold=True)
        else:
            run = paragraph.add_run(part)
            if base_size is not None:
                run.font.size = Pt(base_size)


def clean_markdown_text(text: str) -> str:
    text = text.replace("\\", "")
    text = text.replace("&nbsp;", " ")
    return text.strip()


def add_body_paragraph(doc: Document, text: str) -> None:
    text = clean_markdown_text(text)
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs_from_markdown(p, text)


def column_widths(headers: list[str]) -> list[int]:
    n = len(headers)
    if n == 1:
        return [9360]
    if n == 2:
        if max(len(h) for h in headers) <= 18:
            return [2600, 6760]
        return [3800, 5560]
    if n == 3:
        return [2100, 3000, 4260]
    if n == 4:
        return [1600, 1900, 2700, 3160]
    widths = [max(1200, int(9360 / n)) for _ in headers]
    widths[-1] += 9360 - sum(widths)
    return widths


def parse_table(block: list[str]) -> tuple[list[str], list[list[str]]]:
    def split_row(row: str) -> list[str]:
        row = row.strip().strip("|")
        return [clean_markdown_text(cell) for cell in row.split("|")]

    headers = split_row(block[0])
    rows = [split_row(row) for row in block[2:]]
    rows = [row + [""] * (len(headers) - len(row)) for row in rows]
    return headers, rows


def add_markdown_table(doc: Document, block: list[str]) -> None:
    headers, rows = parse_table(block)
    if not headers:
        return
    widths = column_widths(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_width(table, 9360)
    set_table_borders(table)
    set_row_repeat_as_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    wide_table = len(headers) >= 7
    table_font_size = 7.2 if wide_table else 8.6
    cell_margin = 55 if wide_table else 120
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell, start=cell_margin, end=cell_margin)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.style = doc.styles["AETHER Table Text"]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.keep_with_next = True
        add_runs_from_markdown(p, headers[idx], base_size=table_font_size)
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = NAVY
    for row in rows:
        table_row = table.add_row()
        set_row_cant_split(table_row)
        cells = table_row.cells
        for idx, cell in enumerate(cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell, start=cell_margin, end=cell_margin)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles["AETHER Table Text"]
            p.paragraph_format.space_after = Pt(0)
            add_runs_from_markdown(p, row[idx] if idx < len(row) else "", base_size=table_font_size)
    keep_compact_table_together(table, len(rows))
    if len(rows) > 1:
        for cell in table.rows[-2].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    doc.add_paragraph()


def image_size_inches(path: Path, max_width: float = 6.1, max_height: float = 6.8) -> tuple[float, float] | None:
    if not path.exists() or Image is None:
        return None
    with Image.open(path) as img:
        width_px, height_px = img.size
    ratio = min(max_width / width_px, max_height / height_px) * 96
    width_in = width_px * ratio / 96
    height_in = height_px * ratio / 96
    return width_in, height_in


def add_image(doc: Document, alt: str, rel: str, figure_number: int) -> bool:
    base = PAPER.parent if rel.startswith("..") else ROOT
    path = resolve_within(base, rel, ROOT)
    if not path.exists():
        file_name = Path(rel).name
        path = resolve_within(FIGURE_DIR, file_name, FIGURE_DIR)
    if not path.exists():
        p = doc.add_paragraph(style="AETHER Caption")
        p.add_run(f"[Missing figure: {rel}]")
        return False

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    size = image_size_inches(path)
    if size is None:
        run = p.add_run()
        run.add_picture(str(path), width=Inches(6.1))
    else:
        run = p.add_run()
        run.add_picture(str(path), width=Inches(size[0]), height=Inches(size[1]))

    caption = doc.add_paragraph(style="AETHER Caption")
    add_runs_from_markdown(caption, f"Figure {figure_number}. {alt.strip() or path.stem.replace('_', ' ')}")
    return True


def make_docx() -> dict[str, int]:
    text = PAPER.read_text(encoding="utf-8")
    text = re.sub(r"Status: Working paper v0\.\d+", f"Status: Working paper {VERSION}", text)
    lines = text.splitlines()

    doc = Document()
    configure_document(doc)
    set_header_footer(doc)
    add_cover(doc)

    figure_count = 0
    table_count = 0
    paragraph_buffer: list[str] = []
    i = 0
    skip_front_matter = True

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_body_paragraph(doc, " ".join(paragraph_buffer))
            paragraph_buffer = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if skip_front_matter:
            if line.startswith("## Abstract"):
                skip_front_matter = False
            else:
                i += 1
                continue

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        if line.strip().startswith("<!--") and line.strip().endswith("-->"):
            flush_paragraph()
            i += 1
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            flush_paragraph()
            figure_count += 1
            add_image(doc, image_match.group(1), image_match.group(2), figure_count)
            i += 1
            continue

        if line.startswith("|"):
            flush_paragraph()
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            if len(block) >= 2:
                add_markdown_table(doc, block)
                table_count += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            heading_text = clean_markdown_text(heading_match.group(2))
            style = "Heading 1" if level <= 2 else "Heading 2" if level == 3 else "Heading 3"
            p = doc.add_paragraph(style=style)
            add_runs_from_markdown(p, heading_text)
            i += 1
            continue

        bullet = re.match(r"^-\s+(.*)", line.strip())
        if bullet:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_runs_from_markdown(p, clean_markdown_text(bullet.group(1)))
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        if numbered:
            flush_paragraph()
            # Preserve the source number literally. Word's built-in List Number
            # style otherwise continues numbering across unrelated lists.
            p = doc.add_paragraph(style="Body Text")
            p.paragraph_format.left_indent = Inches(0.28)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.add_run(f"{numbered.group(1)}. ")
            add_runs_from_markdown(p, clean_markdown_text(numbered.group(2)))
            i += 1
            continue

        paragraph_buffer.append(line.strip())
        i += 1

    flush_paragraph()

    props = doc.core_properties
    props.title = "AETHER: Atmospheric Engineering Through High-Energy Removal"
    props.subject = "A conditional feasibility analysis for 100 GtCO2/year removal in an AI- and robotics-accelerated economy"
    props.author = "Noah Hicks"
    props.comments = f"Generated from the AETHER {VERSION} Markdown manuscript by scripts/build_aether_word_manuscript.py."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return {
        "figures": figure_count,
        "tables": table_count,
        "paragraphs": len(doc.paragraphs),
        "sections": len(doc.sections),
        "inline_shapes": len(doc.inline_shapes),
    }


def validate_docx(path: Path, expected_figures: int) -> dict[str, int | str]:
    if not path.exists() or path.stat().st_size == 0:
        raise RuntimeError(f"DOCX was not created: {path}")
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        media_count = len([name for name in names if name.startswith("word/media/")])
        if "word/document.xml" not in names:
            raise RuntimeError("DOCX missing word/document.xml")
    doc = Document(path)
    if len(doc.inline_shapes) < expected_figures:
        raise RuntimeError(f"Expected at least {expected_figures} inline figures, found {len(doc.inline_shapes)}")
    return {
        "size_kb": round(path.stat().st_size / 1024),
        "media_count": media_count,
        "inline_shapes": len(doc.inline_shapes),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
    }


def legacy_main() -> None:
    stats = make_docx()
    validation = validate_docx(OUTPUT, stats["figures"])

    manifest = f"""# AETHER Word Export Manifest

Last updated: 2026-08-10

Generated DOCX:

- `manuscript/submission/{DOCX_NAME}`

## Build Inputs

- Source manuscript: `manuscript/paper/aether_scientific_paper.md`
- Design preset: `narrative_proposal`
- Header pattern: `proposal_centerpiece` adapted for a serious academic working paper
- Builder: `scripts/build_aether_word_manuscript.py`

## Structural Checks

- Markdown figures processed: {stats["figures"]}
- DOCX inline shapes: {validation["inline_shapes"]}
- Embedded media files: {validation["media_count"]}
- Markdown tables processed: {stats["tables"]}
- DOCX tables: {validation["tables"]}
- Paragraphs: {validation["paragraphs"]}
- File size: {validation["size_kb"]} KB

## Notes

This is a professional Word review copy of the AETHER {VERSION} working paper. It preserves the paper's caution that AETHER is not yet publication-ready: climate-model and species-emissions gates remain unresolved.
"""
    MANIFEST.write_text(manifest, encoding="utf-8")

    notes = f"""# Word Export and DOCX Readiness Notes

Last updated: 2026-08-10

The Word export creates `manuscript/submission/{DOCX_NAME}` from the {VERSION} AETHER working paper. It uses a restrained proposal style: formal cover, running header/footer, styled headings, fixed-width tables, rendered figures, and source-preserving references.

The builder processed {stats["figures"]} figures and {stats["tables"]} Markdown tables. Structural validation found {validation["inline_shapes"]} inline image shapes and {validation["media_count"]} embedded media files.

The DOCX is intended for Word-based academic review. It is not a journal-formatted final submission until a target venue is chosen and the citation/figure/table rules are applied.
"""
    NOTES.write_text(notes, encoding="utf-8")

    print(f"Wrote {OUTPUT}")
    print(f"Wrote {MANIFEST}")
    print(f"Wrote {NOTES}")
    print(f"AETHER Word export complete: {validation['inline_shapes']} figures, {validation['tables']} tables, {validation['size_kb']} KB.")


if __name__ == "__main__":
    legacy_main()

